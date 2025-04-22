from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from collections import Counter

def export_to_docx(publicacoes, duplicados, buffer):
    doc = Document()
    titulo = doc.add_heading("Diário Consolidado", 0)
    titulo.runs[0].font.color.rgb = RGBColor(15, 30, 68)

    doc.add_paragraph(f"Data de geração: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    doc.add_paragraph(f"Total de publicações nesta parte: {len(publicacoes)}")
    doc.add_paragraph(f" - Duplicados: {len(duplicados)}")

    origens = [origem for pub in publicacoes for origem in pub.get("origens", [pub["origem"]])]
    contagem = Counter(origens)
    doc.add_paragraph("Publicações por arquivo:")
    for nome, qtd in contagem.items():
        doc.add_paragraph(f" - {nome}: {qtd}")
    doc.add_paragraph("")

    total = len(publicacoes)
    for i, pub in enumerate(publicacoes, 1):
        table = doc.add_table(rows=1, cols=1)
        table.allow_autofit = True
        cell = table.cell(0, 0)

        run = cell.paragraphs[0].add_run(f"Publicação {i} de {total}\n")
        run.bold = True

        for linha in pub["texto"].splitlines():
            if linha.strip():
                p = cell.add_paragraph(linha.strip())
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        if "duplicado_de" in pub:
            cell.add_paragraph("\nDuplicada também encontrada em:")
            for origem_info in pub["duplicado_de"]:
                cell.add_paragraph(f" - {origem_info}")

        doc.add_paragraph("")

    doc.add_paragraph("________________________________________")
    doc.add_paragraph("Criado por Maria Clara Nogueira Diniz - OAB/PI 23765")
    doc.add_paragraph("Site: www.mcts.adv.br | E-mail: contato@mcts.adv.br")
    doc.save(buffer)
    buffer.seek(0)